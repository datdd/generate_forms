import os
import sys
import logging
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# Initialize logging configuration
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# --- Tunables ---
EXCEL_FILE_PATH = "data/input.xlsx"
TEMPLATE_PATH = "data/template.docx"
OUTPUT_DIR = "output"
SHEET_NAME = 0  

# Font Configuration
TARGET_FONT_NAME = "Times New Roman"
TARGET_FONT_SIZE_PT = 12

FIELD_MAPPING = {
    "<<FULL_NAME>>": "Họ và tên",
    "<<CLASS_NAME>>": "Lớp"
}

SUCCESS_CODE = 0
ERROR_MISSING_FILE = 1
ERROR_PROCESSING = 2
# ----------------

def sanitize_filename(filename: str) -> str:
    """
    Remove invalid characters from a string to make it safe for file systems.
    """
    if not isinstance(filename, str):
        return "Unknown"
    safe_name = re.sub(r'[^\w\s-]', '', filename)
    return safe_name.strip()

def apply_font_settings(paragraph, font_name: str, font_size_pt: int) -> None:
    """
    Enforce font family and font size on all runs within a paragraph.
    Handles XML core attributes for robust rendering of Unicode/Vietnamese chars.
    """
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        
        # Explicitly modify low-level XML elements to prevent Word from falling back
        # to default fonts (like Calibri) for Vietnamese character blocks.
        r_pr = run._r.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn('w:ascii'), font_name)
        r_fonts.set(qn('w:hAnsi'), font_name)
        r_fonts.set(qn('w:eastAsia'), font_name)
        r_fonts.set(qn('w:cs'), font_name)

def replace_text_in_paragraphs(paragraphs: list, mapping: dict, row_data: pd.Series) -> None:
    """
    Iterate through paragraphs, replace placeholder text, and enforce font formatting.
    """
    for paragraph in paragraphs:
        is_modified = False
        for placeholder, column_name in mapping.items():
            if placeholder in paragraph.text:
                raw_value = row_data.get(column_name)
                str_value = str(raw_value).strip() if pd.notna(raw_value) else "N/A"
                
                paragraph.text = paragraph.text.replace(placeholder, str_value)
                is_modified = True
        
        # Apply font styles only if the paragraph text was changed
        if is_modified:
            apply_font_settings(paragraph, TARGET_FONT_NAME, TARGET_FONT_SIZE_PT)

def process_documents() -> int:
    """
    Main execution loop. Reads Excel, iterates rows, and generates stylized Word documents.
    """
    if not os.path.exists(EXCEL_FILE_PATH):
        logging.error(f"Missing input Excel file: {EXCEL_FILE_PATH}")
        return ERROR_MISSING_FILE
        
    if not os.path.exists(TEMPLATE_PATH):
        logging.error(f"Missing Word template file: {TEMPLATE_PATH}")
        return ERROR_MISSING_FILE

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    try:
        df = pd.read_excel(
            EXCEL_FILE_PATH, 
            sheet_name=SHEET_NAME, 
            engine='openpyxl'
        )
    except Exception as e:
        logging.error(f"Failed to read Excel file: {e}")
        return ERROR_PROCESSING

    target_column = FIELD_MAPPING.get("<<FULL_NAME>>")
    if target_column not in df.columns:
         logging.error(f"Column '{target_column}' not found in Excel sheet.")
         return ERROR_PROCESSING

    df_valid = df.dropna(subset=[target_column])
    
    success_count = 0
    total_count = len(df_valid)

    for index, row in df_valid.iterrows():
        student_name = str(row[target_column]).strip()
        if not student_name or student_name == "nan":
            continue
            
        safe_name = sanitize_filename(student_name)
        output_filename = f"Phieu_Cham_Diem_{safe_name}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        try:
            doc = Document(TEMPLATE_PATH)
            
            # Process standard paragraphs
            replace_text_in_paragraphs(doc.paragraphs, FIELD_MAPPING, row)
            
            # Process paragraphs inside tables
            for table in doc.tables:
                for table_row in table.rows:
                    for cell in table_row.cells:
                        replace_text_in_paragraphs(cell.paragraphs, FIELD_MAPPING, row)
            
            doc.save(output_path)
            logging.info(f"Generated: {output_filename}")
            success_count += 1
            
        except Exception as e:
            logging.error(f"Failed to process document for '{student_name}': {e}")
            continue

    logging.info(f"Processing complete. Success: {success_count}/{total_count}")
    return SUCCESS_CODE

if __name__ == "__main__":
    sys.exit(process_documents())